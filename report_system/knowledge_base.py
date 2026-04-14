# 知识库接入模块
from haystack import Pipeline, Document
from haystack.components.retrievers import InMemoryBM25Retriever
from haystack.components.rankers import TransformersSimilarityRanker
from haystack.components.readers import ExtractiveReader
from haystack.document_stores.in_memory import InMemoryDocumentStore
from haystack.utils import ComponentDevice
import fitz  # PyMuPDF
import docx
import markdown
import requests
from bs4 import BeautifulSoup

class MockTextEmbedder:
    def run(self, text):
        return {"embedding": [0.0] * 768}

class MockDocumentEmbedder:
    def run(self, documents):
        return {"documents": documents}

class KnowledgeBase:
    def __init__(self):
        self.document_store = InMemoryDocumentStore()
        self.pipeline = None
    
    def add_document(self, file_path, doc_type):
        """添加文档到知识库"""
        text = self._extract_text(file_path, doc_type)
        if text:
            from haystack import Document
            document = Document(content=text, meta={"file_path": file_path, "doc_type": doc_type})
            self.document_store.write_documents([document])
            return True
        return False
    
    def _extract_text(self, file_path, doc_type):
        """从不同类型的文件中提取文本"""
        if doc_type == "pdf":
            return self._extract_pdf(file_path)
        elif doc_type == "docx":
            return self._extract_docx(file_path)
        elif doc_type == "markdown":
            return self._extract_markdown(file_path)
        elif doc_type == "txt":
            return self._extract_txt(file_path)
        elif doc_type == "web":
            return self._extract_web(file_path)
        else:
            return None
    
    def _extract_pdf(self, file_path):
        """从PDF中提取文本"""
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        return text
    
    def _extract_docx(self, file_path):
        """从Word文档中提取文本"""
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    
    def _extract_markdown(self, file_path):
        """从Markdown文件中提取文本"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        # 转换为纯文本
        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text()
    
    def _extract_txt(self, file_path):
        """从文本文件中提取文本"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _extract_web(self, url):
        """从网页中提取文本"""
        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')
        # 移除脚本和样式
        for script in soup(['script', 'style']):
            script.decompose()
        return soup.get_text()
    
    def build_pipeline(self):
        """构建检索 pipeline"""
        # 创建组件
        bm25_retriever = InMemoryBM25Retriever(document_store=self.document_store)
        
        # 构建 pipeline
        self.pipeline = Pipeline()
        self.pipeline.add_component("bm25_retriever", bm25_retriever)
    
    def retrieve(self, query, top_k=5):
        """检索相关文档"""
        if not self.pipeline:
            self.build_pipeline()
        
        result = self.pipeline.run(
            {
                "bm25_retriever": {"query": query, "top_k": top_k}
            }
        )
        return result["bm25_retriever"]["documents"]
