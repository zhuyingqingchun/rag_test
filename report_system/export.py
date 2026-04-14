# 导出模块
import os
import markdown
from docx import Document
from weasyprint import HTML

class Exporter:
    def __init__(self):
        pass
    
    def export_markdown(self, content, output_path):
        """导出为Markdown文件"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return output_path
    
    def export_docx(self, content, output_path):
        """导出为Word文档"""
        doc = Document()
        
        # 解析Markdown内容
        lines = content.split('\n')
        current_heading_level = 0
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_paragraph:
                    current_paragraph.add_run('\n')
                continue
            
            # 处理标题
            if line.startswith('#'):
                heading_level = line.count('#')
                heading_text = line.strip('#').strip()
                
                if heading_level == 1:
                    doc.add_heading(heading_text, level=0)
                else:
                    doc.add_heading(heading_text, level=heading_level-1)
                current_paragraph = None
            # 处理列表
            elif line.startswith('- '):
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                current_paragraph.add_run(line[2:]).bold = False
            # 处理普通文本
            else:
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                current_paragraph.add_run(line + ' ')
        
        doc.save(output_path)
        return output_path
    
    def export_pdf(self, content, output_path):
        """导出为PDF文件"""
        # 转换Markdown为HTML
        html = markdown.markdown(content)
        
        # 添加基本样式
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }}
                h1 {{ color: #333; }}
                h2 {{ color: #555; }}
                h3 {{ color: #777; }}
                ul {{ margin-left: 20px; }}
                li {{ margin-bottom: 5px; }}
            </style>
        </head>
        <body>
            {html}
        </body>
        </html>
        """
        
        # 使用weasyprint生成PDF
        HTML(string=html_content).write_pdf(output_path)
        return output_path
    
    def export(self, content, output_path, format):
        """根据格式导出报告"""
        if format == "markdown":
            return self.export_markdown(content, output_path)
        elif format == "docx":
            return self.export_docx(content, output_path)
        elif format == "pdf":
            return self.export_pdf(content, output_path)
        else:
            raise ValueError(f"Unknown format: {format}")
